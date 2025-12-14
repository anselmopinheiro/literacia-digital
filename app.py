#v24
import os
import json
from docx import Document
from datetime import datetime
from docx.oxml.shared import OxmlElement
from docx.oxml.ns import qn

def add_hyperlink(paragraph, url, text):
    """
    Adiciona um hyperlink a um parágrafo
    
    Args:
        paragraph: Parágrafo do documento
        url: URL do link
        text: Texto a ser exibido
    """
    # Cria o elemento hyperlink
    part = paragraph.part
    r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)
    
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    
    # Cria um novo run com o texto
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    
    # Estilo de hyperlink (azul e sublinhado)
    color = OxmlElement('w:color')
    color.set(qn('w:val'), '0563C1')  # Azul padrão
    rPr.append(color)
    
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)
    
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)
    
    return hyperlink

def substituir_tags_docx(template_path, dados, output_path):
    """
    Substitui as tags no documento DOCX pelos valores fornecidos preservando a formatação
    
    Args:
        template_path (str): Caminho para o ficheiro template
        dados (dict): Dicionário com os valores para substituir as tags
        output_path (str): Caminho para o ficheiro de saída
    """
    # Carrega o documento template
    doc = Document(template_path)
    
    # Função para substituir texto em runs preservando formatação
    def substituir_em_runs(runs, tag, valor):
        """Substitui tag em runs preservando formatação"""
        # Junta todo o texto dos runs para verificar se contém a tag
        texto_completo = ''.join(run.text for run in runs)
        
        if tag in texto_completo:
            # Encontra a posição da tag no texto completo
            inicio_tag = texto_completo.find(tag)
            fim_tag = inicio_tag + len(tag)
            
            # Percorre os runs para encontrar onde a tag está
            pos_atual = 0
            run_inicio = None
            run_fim = None
            pos_inicio_no_run = 0
            pos_fim_no_run = 0
            
            for i, run in enumerate(runs):
                texto_run = run.text
                
                # Verifica se o início da tag está neste run
                if run_inicio is None and pos_atual <= inicio_tag < pos_atual + len(texto_run):
                    run_inicio = i
                    pos_inicio_no_run = inicio_tag - pos_atual
                
                # Verifica se o fim da tag está neste run
                if pos_atual < fim_tag <= pos_atual + len(texto_run):
                    run_fim = i
                    pos_fim_no_run = fim_tag - pos_atual
                    break
                    
                pos_atual += len(texto_run)
            
            if run_inicio is not None and run_fim is not None:
                # Caso 1: Tag está inteiramente dentro de um único run
                if run_inicio == run_fim:
                    run = runs[run_inicio]
                    texto_anterior = run.text[:pos_inicio_no_run]
                    texto_posterior = run.text[pos_fim_no_run:]
                    run.text = texto_anterior + str(valor) + texto_posterior
                
                # Caso 2: Tag está espalhada por múltiplos runs
                else:
                    # Modifica o run do início
                    runs[run_inicio].text = runs[run_inicio].text[:pos_inicio_no_run] + str(valor)
                    
                    # Remove texto dos runs intermediários
                    for i in range(run_inicio + 1, run_fim):
                        runs[i].text = ''
                    
                    # Modifica o run do fim
                    runs[run_fim].text = runs[run_fim].text[pos_fim_no_run:]
    
    # Função para substituir tag padlet com hyperlink
    def substituir_padlet_com_link(paragraph, tag, url):
        """Substitui tag <<padlet>> com hyperlink"""
        if tag in paragraph.text:
            # Encontra a posição da tag
            texto = paragraph.text
            inicio_tag = texto.find(tag)
            
            # Separa texto antes e depois da tag
            texto_antes = texto[:inicio_tag]
            texto_depois = texto[inicio_tag + len(tag):]
            
            # Remove https:// do texto visível
            texto_visivel = url.replace("https://", "").replace("http://", "")
            
            # Limpa o parágrafo
            paragraph.clear()
            
            # Adiciona texto antes (se existir)
            if texto_antes:
                paragraph.add_run(texto_antes)
            
            # Adiciona o hyperlink (URL completo, mas texto sem https://)
            hyperlink = add_hyperlink(paragraph, url, texto_visivel)
            paragraph._element.append(hyperlink)
            
            # Adiciona texto depois (se existir)
            if texto_depois:
                paragraph.add_run(texto_depois)
            
            return True
        return False

    # Substitui as tags nos parágrafos
    for paragraph in doc.paragraphs:
        # Trata tag <<padlet>> especialmente com hyperlink
        if "<<padlet>>" in paragraph.text and "<<padlet>>" in dados:
            substituir_padlet_com_link(paragraph, "<<padlet>>", dados["<<padlet>>"])
        else:
            # Outras tags normalmente
            for tag, valor in dados.items():
                if tag != "<<padlet>>" and tag in paragraph.text:
                    substituir_em_runs(paragraph.runs, tag, valor)
    
    # Substitui as tags nas tabelas (se existirem)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    # Trata tag <<padlet>> especialmente com hyperlink
                    if "<<padlet>>" in paragraph.text and "<<padlet>>" in dados:
                        substituir_padlet_com_link(paragraph, "<<padlet>>", dados["<<padlet>>"])
                    else:
                        # Outras tags normalmente
                        for tag, valor in dados.items():
                            if tag != "<<padlet>>" and tag in paragraph.text:
                                substituir_em_runs(paragraph.runs, tag, valor)
    
    # Substitui as tags nos headers e footers
    for section in doc.sections:
        # Header
        if section.header:
            for paragraph in section.header.paragraphs:
                # Trata tag <<padlet>> especialmente com hyperlink
                if "<<padlet>>" in paragraph.text and "<<padlet>>" in dados:
                    substituir_padlet_com_link(paragraph, "<<padlet>>", dados["<<padlet>>"])
                else:
                    # Outras tags normalmente
                    for tag, valor in dados.items():
                        if tag != "<<padlet>>" and tag in paragraph.text:
                            substituir_em_runs(paragraph.runs, tag, valor)
        
        # Footer
        if section.footer:
            for paragraph in section.footer.paragraphs:
                # Trata tag <<padlet>> especialmente com hyperlink
                if "<<padlet>>" in paragraph.text and "<<padlet>>" in dados:
                    substituir_padlet_com_link(paragraph, "<<padlet>>", dados["<<padlet>>"])
                else:
                    # Outras tags normalmente
                    for tag, valor in dados.items():
                        if tag != "<<padlet>>" and tag in paragraph.text:
                            substituir_em_runs(paragraph.runs, tag, valor)
    
    # Guarda o documento
    doc.save(output_path)
    print(f"Ficheiro criado: {output_path}")

def criar_json_configuracao():
    """
    Cria um ficheiro JSON de exemplo com a configuração das 3 turmas
    """
    configuracao = {
        "turmas": {
            "TurmaA": {
                "nome": "TurmaA",
                "dt": "Direção Técnica A",
                "ronda": "1ª Ronda",
                "padlet": "https://padlet.com/turmaA",
                "docentes": {
                    "docente1": "Prof. João Silva",
                    "docente2": "Prof. Maria Santos", 
                    "docente3": "Prof. Carlos Pereira",
                    "docente4": "Prof. Ana Costa"
                },
                "sessoes": [
                    {"sessao": 1, "data": "2024-01-15"},
                    {"sessao": 2, "data": "2024-01-22"},
                    {"sessao": 3, "data": "2024-01-29"},
                    {"sessao": 4, "data": "2024-02-05"}
                ]
            },
            "TurmaB": {
                "nome": "TurmaB",
                "dt": "Direção Técnica B", 
                "ronda": "2ª Ronda",
                "padlet": "https://padlet.com/turmaB",
                "docentes": {
                    "docente1": "Prof. Ricardo Oliveira",
                    "docente2": "Prof. Luísa Fernandes",
                    "docente3": "Prof. Miguel Torres",
                    "docente4": "Prof. Sofia Ribeiro"
                },
                "sessoes": [
                    {"sessao": 1, "data": "2024-02-12"},
                    {"sessao": 2, "data": "2024-02-19"},
                    {"sessao": 3, "data": "2024-02-26"},
                    {"sessao": 4, "data": "2024-03-05"},
                    {"sessao": 5, "data": "2024-03-12"}
                ]
            },
            "TurmaC": {
                "nome": "TurmaC",
                "dt": "Direção Técnica C",
                "ronda": "3ª Ronda", 
                "padlet": "https://padlet.com/turmaC",
                "docentes": {
                    "docente1": "Prof. Pedro Almeida",
                    "docente2": "Prof. Carla Mendes",
                    "docente3": "Prof. Rui Teixeira", 
                    "docente4": "Prof. Isabel Rocha"
                },
                "sessoes": [
                    {"sessao": 1, "data": "2024-03-19"},
                    {"sessao": 2, "data": "2024-03-26"},
                    {"sessao": 3, "data": "2024-04-02"}
                ]
            }
        }
    }
    
    with open('configuracao_turmas.json', 'w', encoding='utf-8') as f:
        json.dump(configuracao, f, indent=4, ensure_ascii=False)
    
    print("Ficheiro 'configuracao_turmas.json' criado com sucesso!")
    print("Pode editar este ficheiro para personalizar as turmas, datas e docentes.")

def carregar_configuracao(json_file):
    """Carrega o ficheiro de configuração JSON caso exista."""
    if not os.path.exists(json_file):
        print(f"ERRO: Ficheiro '{json_file}' não encontrado!")
        print("Execute a opção 1 primeiro para criar o ficheiro de configuração.")
        return None

    try:
        with open(json_file, 'r', encoding='utf-8') as f:
            return json.load(f)
    except Exception as e:
        print(f"ERRO ao ler ficheiro JSON: {e}")
        return None

def listar_turmas_disponiveis(config):
    """Mostra na consola as turmas disponíveis no ficheiro de configuração."""
    turmas = list(config.get("turmas", {}).keys())

    if not turmas:
        print("Não foram encontradas turmas no ficheiro JSON.")
        return []

    print("=== TURMAS DISPONÍVEIS ===")
    for idx, turma in enumerate(turmas, start=1):
        print(f"{idx}. {turma}")

    return turmas

def processar_turmas_do_json(turmas_selecionadas=None, config=None):
    """Processa todas as turmas ou apenas as selecionadas no ficheiro JSON."""
    json_file = "configuracao_turmas.json"
    template_file = "template.docx"
    output_folder = "documentos_turmas"

    if not os.path.exists(template_file):
        print(f"ERRO: Ficheiro template '{template_file}' não encontrado!")
        return

    if config is None:
        config = carregar_configuracao(json_file)
        if config is None:
            return

    turmas_config = config.get("turmas", {})
    if not turmas_config:
        print("ERRO: Não foram encontradas turmas no ficheiro de configuração.")
        return

    if turmas_selecionadas:
        turmas_config = {nome: dados for nome, dados in turmas_config.items() if nome in turmas_selecionadas}
        if not turmas_config:
            print("ERRO: Nenhuma das turmas selecionadas existe no ficheiro JSON.")
            return

    if not os.path.exists(output_folder):
        os.makedirs(output_folder)

    turmas_processadas = 0

    for nome_turma, dados_turma in turmas_config.items():
        turmas_processadas += 1
        print(f"\n=== Processando {nome_turma} ===")

        for sessao_info in dados_turma["sessoes"]:
            try:
                data_obj = datetime.strptime(sessao_info["data"], "%Y-%m-%d")
                data_formatada = data_obj.strftime("%d/%m/%Y")

                dados_documento = {
                    "<<Turma>>": dados_turma["nome"],
                    "<<DT>>": dados_turma["dt"],
                    "<<ronda>>": dados_turma["ronda"],
                    "<<padlet>>": dados_turma["padlet"],
                    "<<sessao>>": str(sessao_info["sessao"]),
                    "<<data>>": data_formatada,
                    "<<Docente1>>": dados_turma["docentes"]["docente1"],
                    "<<Docente2>>": dados_turma["docentes"]["docente2"],
                    "<<Docente3>>": dados_turma["docentes"]["docente3"],
                    "<<Docente4>>": dados_turma["docentes"]["docente4"]
                }

                ronda_numero = ''.join(filter(str.isdigit, dados_turma["ronda"]))
                if not ronda_numero:
                    ronda_numero = "1"

                nome_arquivo = f"{dados_turma['nome']}_sessao-{sessao_info['sessao']}_ronda-{ronda_numero}_{data_obj.strftime('%Y-%m-%d')}.docx"
                caminho_saida = os.path.join(output_folder, nome_arquivo)

                substituir_tags_docx(template_file, dados_documento, caminho_saida)
                print(f"  → Sessão {sessao_info['sessao']:02d} - {data_formatada} → {nome_arquivo}")

            except ValueError as e:
                print(f"  ERRO: Data inválida '{sessao_info['data']}': {e}")
            except Exception as e:
                print(f"  ERRO ao processar sessão: {e}")

    if turmas_processadas:
        print(f"\n✅ Processamento concluído!")
        print("📁 Ficheiros criados na pasta '{}'".format(output_folder))
        print("📋 Padrão de nomes: turma_sessao-#_ronda-#_data.docx")
        print("   - Turma: nome da turma do JSON")
        print("   - #: número da sessão")
        print("   - #: número da ronda")
        print("   - Data: formato YYYY-MM-DD")

def selecionar_turma_para_processar():
    """Lista as turmas disponíveis e permite processar apenas uma delas."""
    json_file = "configuracao_turmas.json"
    config = carregar_configuracao(json_file)

    if config is None:
        return

    turmas = listar_turmas_disponiveis(config)
    if not turmas:
        return

    escolha = input("\nIntroduza o número da turma que deseja processar (0 para cancelar): ").strip()

    if not escolha.isdigit():
        print("Opção inválida. Utilize apenas números.")
        return

    escolha_num = int(escolha)

    if escolha_num == 0:
        print("Operação cancelada pelo utilizador.")
        return

    if not 1 <= escolha_num <= len(turmas):
        print("Opção inválida. Selecione um número da lista apresentada.")
        return

    turma_escolhida = turmas[escolha_num - 1]
    print(f"\n→ A processar apenas a {turma_escolhida}...")
    processar_turmas_do_json({turma_escolhida}, config=config)

def mostrar_estrutura_json():
    """
    Mostra exemplo da estrutura do JSON para referência
    """
    print("=== ESTRUTURA DO FICHEIRO JSON ===")
    print("""
O ficheiro 'configuracao_turmas.json' tem esta estrutura:

{
    "turmas": {
        "TurmaA": {
            "nome": "TurmaA",
            "dt": "Direção Técnica A",
            "ronda": "1ª Ronda",
            "padlet": "https://padlet.com/turmaA",
            "docentes": {
                "docente1": "Prof. João Silva",
                "docente2": "Prof. Maria Santos",
                "docente3": "Prof. Carlos Pereira",
                "docente4": "Prof. Ana Costa"
            },
            "sessoes": [
                {"sessao": 1, "data": "2024-01-15"},
                {"sessao": 2, "data": "2024-01-22"}
            ]
        }
    }
}

TAGS que serão substituídas no template:
- <<Turma>> → nome da turma
- <<DT>> → direção técnica
- <<ronda>> → ronda
- <<padlet>> → URL do padlet (convertido em HYPERLINK clicável)
- <<sessao>> → número da sessão
- <<data>> → data da sessão (DD/MM/YYYY)
- <<Docente1>> a <<Docente4>> → nomes dos docentes

NOTA: A tag <<padlet>> será convertida num hyperlink clicável azul e sublinhado.

NOME DOS FICHEIROS: turma_sessao-#_ronda-#_data.docx
- turma: nome da turma conforme definido no JSON
- # da sessão: número da sessão (ex: "1", "2", "10")
- # da ronda: número extraído da tag ronda (ex: "1ª Ronda" → "1")
- data: formato YYYY-MM-DD

Exemplos de nomes:
- TurmaA_sessao-1_ronda-1_2024-01-15.docx
- TurmaB_sessao-5_ronda-2_2024-02-12.docx
- TurmaC_sessao-12_ronda-3_2024-03-19.docx
    """)

if __name__ == "__main__":
    print("=== GERADOR DE DOCUMENTOS DOCX COM JSON ===")
    print("1. Criar ficheiro de configuração JSON (configuracao_turmas.json)")
    print("2. Processar todas as turmas do JSON")
    print("3. Listar turmas do JSON e processar uma específica")
    print("4. Mostrar estrutura do JSON")

    opcao = input("\nEscolha uma opção (1-4): ").strip()
    
    if opcao == "1":
        criar_json_configuracao()
        print("\n💡 Dica: Edite o ficheiro 'configuracao_turmas.json' para personalizar:")
        print("   - Nomes das turmas e docentes")
        print("   - Datas das sessões")
        print("   - Direção técnica e rondas")
        print("   - URLs do Padlet (serão convertidos em hyperlinks)")
        
    elif opcao == "2":
        processar_turmas_do_json()

    elif opcao == "3":
        selecionar_turma_para_processar()

    elif opcao == "4":
        mostrar_estrutura_json()
        
    else:
        print("Opção inválida. Tente novamente.")
    
    print("\n📋 Lembre-se:")
    print("   - Tenha o ficheiro 'template.docx' no mesmo diretório")
    print("   - As tags no template devem ser: <<Turma>>, <<DT>>, <<ronda>>, <<sessao>>, <<data>>, <<Docente1>>, <<Docente2>>, <<Docente3>>, <<Docente4>>")
    print("   - A tag <<padlet>> será convertida em hyperlink clicável")
    print("   - Os ficheiros serão criados como: turma_sessao-#_ronda-#_data.docx")